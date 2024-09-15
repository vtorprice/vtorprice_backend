import datetime
import os
from docx import Document
from rest_framework import serializers

from config import settings
from finance.models import InvoicePayment
from logistics.models import TransportApplication
from exchange.models import RecyclablesDeal, EquipmentApplication


class BaseGenerator:
    replacing_mapping: dict[str, str] = None
    template_file_path: str = None
    document = None
    output_file_name: str = None
    input_template_file_path = None

    def __init__(self):
        if not os.path.exists(f"{settings.MEDIA_ROOT}/generated_storage"):
            os.makedirs(f"{settings.MEDIA_ROOT}/generated_storage")

    def replace_all_and_save(self):
        for template_string in self.replacing_mapping.keys():
            self.replace_string(
                template_string, self.replacing_mapping[template_string]
            )
            self.replace_string_in_table(
                template_string, self.replacing_mapping[template_string]
            )
        return self.save()

    def save(self):
        self.document.save(f"{settings.MEDIA_ROOT}/{self.output_file_name}")
        return self.output_file_name

    def replace_string(self, string_to_replace, replace_to):
        """Replaces given string to another in documents paragraph"""
        for p in self.document.paragraphs:
            if string_to_replace in p.text:
                p.text = p.text.replace(string_to_replace, str(replace_to))
                inline = p.runs
                for i in range(len(inline)):
                    if string_to_replace in inline[i].text:
                        text = inline[i].text.replace(
                            string_to_replace, replace_to
                        )
                        inline[i].text = text

    def replace_string_in_table(self, string_to_replace, replace_to):
        """Replaces given string to another in documents table"""
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if string_to_replace in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                string_to_replace, str(replace_to)
                            )

    def convert_num_to_string(self, num):
        import num2words

        return num2words.num2words(num, lang="ru")


class TransportApplicationGeneratorMixin:
    def __init__(self, transport_application: TransportApplication):
        super().__init__()
        self.transport_application = transport_application

    def build_replacing_mapping(self):
        from company.models import Company

        deal = getattr(self.transport_application, "deal", None)
        if not deal:
            raise serializers.ValidationError(
                "Transport application has no deal"
            )

        supplier_company: Company = deal.supplier_company
        buyer_company: Company = deal.buyer_company
        return {
            "%buyer_company_name%": buyer_company.name,
            "%buyer_company_address%": buyer_company.address,
            "%buyer_company_inn%": buyer_company.inn,
            "%seller_company_name%": supplier_company.name,
            "%seller_company_address%": supplier_company.address,
            "%seller_company_inn%": supplier_company.inn,
        }


class ContractorApplicationGeneratorMixin(TransportApplicationGeneratorMixin):
    def build_replacing_mapping(self):
        approved_offer = self.transport_application.approved_logistics_offer
        if not approved_offer:
            raise serializers.ValidationError(
                "Transport application has no approved offer"
            )
        contractor = approved_offer.contractor

        mappings = super().build_replacing_mapping()
        mappings.update(
            {
                "%contractor_name%": contractor.name if contractor else "",
                "%contractor_address%": contractor.address
                if contractor
                else "",
            }
        )
        return mappings


class ContractorWithDeliveryInfoMixin(ContractorApplicationGeneratorMixin):
    def build_replacing_mapping(self):
        deal = self.transport_application.deal
        approved_offer = self.transport_application.approved_logistics_offer
        mappings = super().build_replacing_mapping()
        mappings.update(
            {
                "%shipping_city_name%": getattr(
                    self.transport_application.shipping_city, "name", ""
                ),
                "%delivery_city_name%": getattr(
                    self.transport_application.delivery_city, "name", ""
                ),
                "%shipping_date%": approved_offer.shipping_date
                if approved_offer
                else "",
                "%shipping_address%": self.transport_application.shipping_address,
                "%delivery_date%": deal.delivery_date if deal else "",
                "%delivery_address%": deal.delivery_address if deal else "",
            }
        )
        return mappings


# Договор на отгрузку
class UnloadingAgreement(ContractorWithDeliveryInfoMixin, BaseGenerator):
    def __init__(self, transport_application: TransportApplication):
        super().__init__(transport_application)
        self.output_file_name = f"generated_storage/Договор на отгрузку по заявке №{self.transport_application.id}.docx"
        self.replacing_mapping = self.build_replacing_mapping()
        self.input_template_file_path = f"{settings.PROJECT_DIR}/document_generator/templates/Copy of Доверенность на отгрузку.docx"
        self.document = Document(self.input_template_file_path)

    def build_replacing_mapping(self):
        return super().build_replacing_mapping()


# ТТН todo: подвязать к сделке
class Waybill(ContractorWithDeliveryInfoMixin, BaseGenerator):
    def __init__(self, transport_application: TransportApplication):
        super().__init__(transport_application)
        self.output_file_name = f"generated_storage/ТТН по заявке №{self.transport_application.id}.docx"
        self.input_template_file_path = f"{settings.PROJECT_DIR}/document_generator/templates/trn-2021.docx"
        self.document = Document(self.input_template_file_path)
        self.replacing_mapping = self.build_replacing_mapping()

    def build_replacing_mapping(self):
        mappings = super().build_replacing_mapping()
        deal = self.transport_application.deal
        mappings.update(
            {
                "%deal_item%": deal.application.recyclables.name
                if deal
                else self.transport_application.cargo_type,
                "%deal_item_volume%": deal.weight
                if deal
                else self.transport_application.weight,
                "%seller_company_phone": str(
                    self.transport_application.deal.supplier_company.phone
                ),
            }
        )
        return mappings


# счет фактура
class Invoice(ContractorWithDeliveryInfoMixin, BaseGenerator):
    def __init__(self, transport_application: TransportApplication):
        super().__init__(transport_application)
        self.output_file_name = f"generated_storage/Счет-фактура по заявке №{self.transport_application.id}.docx"
        self.input_template_file_path = f"{settings.PROJECT_DIR}/document_generator/templates/Счет-фактура.docx"
        self.document = Document(self.input_template_file_path)
        self.replacing_mapping = self.build_replacing_mapping()

    def build_replacing_mapping(self):
        deal = self.transport_application.deal
        mappings = super().build_replacing_mapping()
        mappings.update(
            {
                # TODO: Если будет сделка по оборудлванию, передалть под equipment
                "%deal_item%": deal.application.recyclables.name
                if deal
                else self.transport_application.cargo_type,
                "%deal_item_volume%": deal.weight
                if deal
                else self.transport_application.weight,
                "%deal_item_price%": deal.price if deal else 0.0,
                "%deal_total_price%": str(float(deal.price) * deal.weight)
                if deal
                else 0.0,
            }
        )

        return mappings


# договор приложение спецификация
class AgreementSpecification(BaseGenerator):
    def __init__(self, deal: RecyclablesDeal):
        super().__init__()
        self.input_template_file_path = f"{settings.PROJECT_DIR}/document_generator/templates/Договор_приложение_спецификация.docx"
        self.document = Document(self.input_template_file_path)
        self.deal = deal
        self.output_file_name = f"generated_storage/Договор_приложение_спецификация по заявке {deal.id}.docx"

        self.replacing_mapping = self.build_replacing_mappings()

    def build_replacing_mappings(self):
        deal = self.deal
        application = deal.application
        recyclables = getattr(application, "recyclables", None)
        buyer_company = getattr(deal, "buyer_company", None)
        seller_company = getattr(deal, "supplier_company", None)

        dict = {
            "%deal_item_volume%": deal.weight,
            "%deal_item_price%": deal.price,
            "%deal_total_price%": str(float(deal.price) * deal.weight),
            "%deal_number%": deal.deal_number,
            "%date%": datetime.datetime.now().strftime("%d.%B.%Y"),
        }

        if seller_company:
            dict.update(
                {
                    "%seller_company_name%": seller_company.name,
                    "%seller_company_email%": seller_company.email,
                    "%seller_company_phone%": str(seller_company.phone),
                    "%seller_company_address%": seller_company.address,
                    "%seller_company_inn%": seller_company.inn,
                }
            )

        if buyer_company:
            dict.update(
                {
                    "%buyer_company_name%": buyer_company.name,
                    "%buyer_company_email%": buyer_company.email,
                    "%buyer_company_phone%": str(buyer_company.phone),
                    "%buyer_company_address%": buyer_company.address,
                    "%buyer_company_inn%": buyer_company.inn,
                    "%city%": getattr(buyer_company.city, "name", ""),
                }
            )

        if application:
            if isinstance(application, EquipmentApplication):
                volume = deal.weight or 0.0
                sum = application.price * application.count
            else:
                volume = deal.weight or 0.0
                sum = float(application.price) * volume

            dict.update(
                {
                    "%recyclables_volume%": volume,
                    "%recyclables_price%": application.price,
                    "%recyclables_sum%": sum,
                    "%total_sum%": sum,
                    "%total_sum_string%": self.convert_num_to_string(sum),
                    "%recyclables_photos%": getattr(application, "images", ""),
                }
            )

        if recyclables:
            dict.update(
                {
                    "%deal_item%": recyclables.name,
                    "%recyclables_name%": recyclables.name,
                    "%recyclables_description%": recyclables.description,
                    "%recyclables_category%": recyclables.category,
                    "%recyclables_subcategory%": recyclables.category.parent,
                }
            )

        return dict


# Договор заявка
class AgreementApplication(ContractorApplicationGeneratorMixin, BaseGenerator):
    def __init__(self, transport_application: TransportApplication):
        super().__init__(transport_application)
        self.input_template_file_path = f"{settings.PROJECT_DIR}/document_generator/templates/Договор-Заявка.docx"
        self.document = Document(self.input_template_file_path)
        self.output_file_name = f"generated_storage/Договор-Заявка по заявке №{self.transport_application.id}.docx"

        self.replacing_mapping = self.build_replacing_mapping()

    def build_replacing_mapping(self):

        deal = self.transport_application.deal
        mapping = super().build_replacing_mapping()
        if not self.transport_application.approved_logistics_offer:
            raise serializers.ValidationError(
                "У заявки нет выбранного предложения логиста"
            )
        mapping.update(
            {
                "%deal_item%": deal.application.recyclables.name,
                "%deal_item_volume%": deal.weight,
                "%price%": deal.application.total_price,
                "%payment_term%": deal.payment_term,
                "%loading_type%": self.transport_application.loading_type,
                "%shipping_city_name%": getattr(
                    self.transport_application.shipping_city, "name", ""
                ),
                "%delivery_city_name%": getattr(
                    self.transport_application.delivery_address, "name", ""
                ),
                "%shipping_date%": self.transport_application.approved_logistics_offer.shipping_date,
                "%shipping_address%": self.transport_application.shipping_address,
                "%item_name%": deal.application.recyclables.name,
                "%item_volume%": deal.weight,
                "%delivery_date%": deal.delivery_date,
                "%delivery_address%": self.transport_application.delivery_address,
            }
        )

        return mapping


# УПД todo: подвязать к сделке
class UniformTransferDocument(
    TransportApplicationGeneratorMixin, BaseGenerator
):
    def __init__(self, transport_application: TransportApplication):
        super().__init__(transport_application)
        self.transport_application = transport_application
        self.input_template_file_path = (
            f"{settings.PROJECT_DIR}/document_generator/templates/УПД.docx"
        )
        self.document = Document(self.input_template_file_path)
        self.replacing_mapping = self.build_replacing_mapping()
        self.output_file_name = f"generated_storage/УПД по заявке 3 № {self.transport_application.id}.docx"

    def build_replacing_mapping(self):
        mappings = super().build_replacing_mapping()
        mappings.update(
            {
                "%contractor_name%": self.transport_application.approved_logistics_offer.contractor.name,
                "%deal_item%": self.transport_application.deal.application.recyclables.name,
                "%deal_item_volume%": self.transport_application.deal.weight,
                "%deal_item_price%": self.transport_application.deal.application.price,
                "%deal_total_price%": self.transport_application.deal.application.total_price,
                "%shipping_date%": self.transport_application.approved_logistics_offer.shipping_date,
                "%delivery_date%": self.transport_application.deal.delivery_date,
            }
        )
        return mappings


# Акт
class Act(BaseGenerator):
    def __init__(self, company, deal):
        super().__init__()
        self.input_template_file_path = (
            f"{settings.PROJECT_DIR}/document_generator/templates/Акт.docx"
        )
        self.document = Document(self.input_template_file_path)
        self.company = company
        self.deal: RecyclablesDeal = deal
        self.price_per_kg = 1
        self.replacing_mapping = self.build_replacing_mapping()
        self.output_file_name = f"generated_storage/Акт по заявке №{self.deal.deal_number} для {self.company.name}.docx"

    def build_replacing_mapping(self):
        return {
            "%buyer_company%": self.company.name,
            "%buyer_company_inn%": self.company.inn,
            "%buyer_company_address%": self.company.address,
            "%buyer_company_phone%": self.company.phone,
            "%count%": self.deal.weight,
            "%price_per_kg%": self.price_per_kg,
            "%total_sum%": self.price_per_kg * self.deal.weight,
        }


class InvoiceDocument(AgreementSpecification, BaseGenerator):
    def __init__(self, invoice: InvoicePayment):
        self.invoice = invoice
        self.price_per_kg = 1

        super().__init__(invoice.deal)

        self.input_template_file_path = f"{settings.PROJECT_DIR}/document_generator/templates/Счёт поставка.docx"
        self.document = Document(self.input_template_file_path)
        self.output_file_name = (
            f"generated_storage/Счёт поставка {invoice.id}.docx"
        )

    def build_replacing_mappings(self):
        mappings = super().build_replacing_mappings()
        mappings.update(
            {
                "%invoice_number%": self.invoice.id,
                "%date": str(datetime.datetime.now()),
                "%count%": self.deal.weight,
                "%price_per_kg%": self.price_per_kg,
                "%total_sum%": self.price_per_kg * self.deal.weight,
            }
        )
        return mappings
